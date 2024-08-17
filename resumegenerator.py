import streamlit as st
from io import BytesIO
import docx

# Create a Streamlit app title
st.title("Resume Generator")

# Define a function to generate the resume
def generate_resume(full_name, email, phone_number, current_job_title, work_experience, key_skills, education_background, certifications_awards, desired_job_position, hobbies_interests, references):
    try:
        # Create a new Word document
        doc = docx.Document()

        # Add a heading with the full name
        doc.add_heading(full_name, 0)

        # Add a section for contact information
        doc.add_paragraph(f"Email: {email}")
        doc.add_paragraph(f"Phone Number: {phone_number}")

        # Add a section for current job title
        doc.add_paragraph(f"Current Job Title: {current_job_title}")

        # Add a section for work experience
        doc.add_heading("Work Experience", 1)
        doc.add_paragraph(work_experience)

        # Add a section for key skills
        doc.add_heading("Key Skills", 1)
        doc.add_paragraph(key_skills)

        # Add a section for education background
        doc.add_heading("Education Background", 1)
        doc.add_paragraph(education_background)

        # Add a section for certifications and awards
        doc.add_heading("Certifications and Awards", 1)
        doc.add_paragraph(certifications_awards)

        # Add a section for desired job position
        doc.add_paragraph(f"Desired Job Position: {desired_job_position}")

        # Add a section for hobbies and interests
        doc.add_heading("Hobbies and Interests", 1)
        doc.add_paragraph(hobbies_interests)

        # Add a section for references
        doc.add_heading("References", 1)
        doc.add_paragraph(references)

        # Save the generated resume to a BytesIO object
        resume_buffer = BytesIO()
        doc.save(resume_buffer)

        # Return the generated resume
        return resume_buffer.getvalue()
    except Exception as e:
        st.error("Error generating resume: {}".format(e))
        return None

# Create a form to collect user information
with st.form("resume-form"):
    full_name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone_number = st.text_input("Phone Number")
    current_job_title = st.text_input("Current Job Title")
    work_experience = st.text_area("Work Experience")
    key_skills = st.text_area("Key Skills")
    education_background = st.text_area("Education Background")
    certifications_awards = st.text_area("Certifications and Awards")
    desired_job_position = st.text_input("Desired Job Position")
    hobbies_interests = st.text_area("Hobbies and Interests")
    references = st.text_area("References")
    submitted = st.form_submit_button("Generate Resume")

# Generate the resume when the form is submitted
if submitted:
    if not full_name or not email or not phone_number:
        st.error("Please fill in all required fields")
    else:
        resume_buffer = generate_resume(full_name, email, phone_number, current_job_title, work_experience, key_skills, education_background, certifications_awards, desired_job_position, hobbies_interests, references)
        if resume_buffer:
            st.write("Resume successfully generated!")
            st.download_button(label="Download Generated Resume", data=resume_buffer, file_name="generated_resume.docx")